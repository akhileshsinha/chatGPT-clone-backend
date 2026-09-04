from docx import Document
from docx.shared import Pt


class WordGenerator:

    def generate(
            self,
            title: str,
            sections: list[dict],
            output_path: str,
    ):
        document = Document()

        # Title
        document.add_heading(
            title,
            level=0,
        )

        for section in sections:
            heading = section.get(
                "heading",
                "",
            )

            content = section.get(
                "content",
                "",
            )

            if heading:
                document.add_heading(
                    heading,
                    level=1,
                )

            if content:
                paragraph = document.add_paragraph(
                    content
                )

                for run in paragraph.runs:
                    run.font.size = Pt(11)

            bullets = section.get(
                "bullets",
                [],
            )

            for bullet in bullets:
                paragraph = document.add_paragraph(
                    str(bullet),
                    style="List Bullet",
                )

        document.save(output_path)