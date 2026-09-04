from docx import Document


class WordModifier:

    def modify(
            self,
            file_path: str,
            modifications: list[dict],
            output_path: str,
    ):
        document = Document(file_path)

        for modification in modifications:
            action = modification.get("action")

            if action == "replace_text":
                old_text = str(
                    modification.get(
                        "old_text",
                        "",
                    )
                )

                new_text = str(
                    modification.get(
                        "new_text",
                        "",
                    )
                )

                for paragraph in document.paragraphs:
                    if old_text in paragraph.text:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(
                                    old_text,
                                    new_text,
                                )

            elif action == "update_paragraph":
                paragraph_index = modification.get(
                    "paragraph_index"
                )

                new_text = modification.get(
                    "text",
                    "",
                )

                if (
                        paragraph_index
                        and 1 <= paragraph_index <= len(
                    document.paragraphs
                )
                ):
                    document.paragraphs[
                        paragraph_index - 1
                        ].text = new_text

            elif action == "add_paragraph":
                text = modification.get(
                    "text",
                    "",
                )

                style = modification.get(
                    "style",
                    None,
                )

                if style:
                    document.add_paragraph(
                        text,
                        style=style,
                    )
                else:
                    document.add_paragraph(
                        text
                    )

            elif action == "remove_paragraph":
                paragraph_index = modification.get(
                    "paragraph_index"
                )

                if (
                        paragraph_index
                        and 1 <= paragraph_index <= len(
                    document.paragraphs
                )
                ):
                    paragraph = document.paragraphs[
                        paragraph_index - 1
                        ]

                    paragraph._element.getparent().remove(
                        paragraph._element
                    )

            elif action == "add_heading":
                text = modification.get(
                    "text",
                    "",
                )

                level = modification.get(
                    "level",
                    1,
                )

                document.add_heading(
                    text,
                    level=level,
                )

        document.save(output_path)